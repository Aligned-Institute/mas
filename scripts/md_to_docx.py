import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INPUT = "/Users/warmachine/Documents/PROJECTS/ALI/R-D/Research/csignals/docs/signals-whitepaper.md"
OUTPUT = "/Users/warmachine/Documents/PROJECTS/ALI/R-D/Research/csignals/docs/signals-whitepaper.docx"

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

def set_style(para, size=11, bold=False, color=None, align=None):
    for run in para.runs:
        run.font.size = Pt(size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*color)
    if align:
        para.alignment = align

def add_heading(doc, text, level):
    sizes = {1: (18, True), 2: (15, True), 3: (13, True)}
    size, bold = sizes.get(level, (11, True))
    p = doc.add_paragraph()
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_table_from_md(doc, rows):
    header = [c.strip() for c in rows[0].split('|') if c.strip()]
    data_rows = []
    for row in rows[2:]:
        cells = [c.strip() for c in row.split('|') if c.strip()]
        if cells:
            data_rows.append(cells)
    if not header or not data_rows:
        return
    col_count = max(len(header), max(len(r) for r in data_rows))
    table = doc.add_table(rows=1 + len(data_rows), cols=col_count)
    table.style = 'Table Grid'
    for i, h in enumerate(header):
        if i < col_count:
            cell = table.cell(0, i)
            cell.text = h
            for run in cell.paragraphs[0].runs:
                run.font.bold = True
                run.font.size = Pt(10)
    for ri, row in enumerate(data_rows):
        for ci, val in enumerate(row):
            if ci < col_count:
                cell = table.cell(ri + 1, ci)
                cell.text = val
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()

def add_code_block(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    run = p.add_run('\n'.join(lines))
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(40, 40, 40)

def inline_bold(para, text):
    parts = re.split(r'\*\*(.+?)\*\*', text)
    for i, part in enumerate(parts):
        if not part:
            continue
        run = para.add_run(part)
        run.font.size = Pt(11)
        if i % 2 == 1:
            run.font.bold = True

with open(INPUT, 'r') as f:
    lines = f.readlines()

i = 0
while i < len(lines):
    line = lines[i].rstrip('\n')

    # H1
    if line.startswith('# ') and not line.startswith('## '):
        add_heading(doc, line[2:], 1)

    # H2
    elif line.startswith('## ') and not line.startswith('### '):
        add_heading(doc, line[3:], 2)

    # H3
    elif line.startswith('### '):
        add_heading(doc, line[4:], 3)

    # HR
    elif re.match(r'^-{3,}$', line):
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'AAAAAA')
        pBdr.append(bottom)
        pPr.append(pBdr)

    # Table
    elif line.startswith('|'):
        table_rows = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            table_rows.append(lines[i].rstrip('\n'))
            i += 1
        add_table_from_md(doc, table_rows)
        continue

    # Code block
    elif line.startswith('```'):
        code_lines = []
        i += 1
        while i < len(lines) and not lines[i].startswith('```'):
            code_lines.append(lines[i].rstrip('\n'))
            i += 1
        add_code_block(doc, code_lines)

    # Bullet
    elif line.startswith('- '):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25)
        inline_bold(p, line[2:])
        for run in p.runs:
            if not run.font.bold:
                run.font.size = Pt(11)

    # Numbered bullet (i. ii. iii.)
    elif re.match(r'^[ivxIVX]+\.\s', line):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)
        inline_bold(p, line)
        for run in p.runs:
            if not run.font.bold:
                run.font.size = Pt(11)

    # Italic block quote style (*text*)
    elif line.startswith('> '):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.right_indent = Inches(0.5)
        run = p.add_run(line[2:])
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = RGBColor(80, 80, 80)

    # Empty line
    elif line.strip() == '':
        pass

    # Normal paragraph
    else:
        p = doc.add_paragraph()
        inline_bold(p, line)
        p.paragraph_format.space_after = Pt(6)

    i += 1

doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
